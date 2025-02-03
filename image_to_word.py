import base64
from io import BytesIO
import json
from PIL import Image
from docx import Document


def decode_base64_to_image(foto_base64):
    # Decodificar e salvar a imagem
    img_data = base64.b64decode(foto_base64)
                
    # Carregar a imagem usando Pillow
    image = Image.open(BytesIO(img_data))
    
    # Definir uma altura fixa, mantendo a proporção
    fixed_height = 200  # altura em pixels
    width_percent = (fixed_height / float(image.size[1]))
    new_width = int((float(image.size[0]) * float(width_percent)))
    
    # Redimensionar a imagem
    resized_image = image.resize((new_width, fixed_height))
    
    # Converter a imagem redimensionada de volta para BytesIO
    image_stream = BytesIO()
    resized_image.save(image_stream, format=image.format)
    image_stream.seek(0)  # Resetar o stream para a leitura
    return image_stream


def insert_images_to_word(json_file_path, docx_file_path):
    """Insere imagens em um local específico no documento Word."""

    # Carregar os dados do JSON
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)  

    # Abrir o documento Word existente
    doc = Document(docx_file_path)

    # Verifica se há a chave 'images' no JSON
    if 'images' not in data:
        print("Erro: O JSON não contém a chave 'images'.")
        return
    image_data = data['images']
    
    
    for para in doc.paragraphs:
        texto = para.text
        print(texto)
        for item in image_data:
            aux = False
            keyword = item['id']
            
            if keyword in texto:
                print("encontrado:", keyword)
                foto = item['fotos']
                if not foto:  # Se 'fotos' for uma lista vazia ou None
                    print(f"Aviso: O item {keyword} tem uma lista vazia de fotos.")
                    continue
                
                image_stream = decode_base64_to_image(foto)
            
                    
                para.text = para.text.replace(keyword, "")
                run = para.add_run()
                run.add_picture(image_stream)
                
                
                  
            else:
                print('key',keyword, 'nao encontrado')
                #break
            
            

    # Salvar o documento atualizado
    doc.save(docx_file_path)
    print(f"Imagem inserida com sucesso no documento {docx_file_path}")


