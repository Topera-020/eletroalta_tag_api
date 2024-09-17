import xml.etree.ElementTree as ET
import shutil
import zipfile
import os
import base64
from docx import Document
from io import BytesIO
from PIL import Image

def convert_json_to_word_xml(raw_json):
    # Definir os namespaces e criar o elemento raiz com os namespaces
    namespaces = {
        'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
        'cx': 'http://schemas.microsoft.com/office/drawing/2014/chartex',
        'cx1': 'http://schemas.microsoft.com/office/drawing/2015/9/8/chartex',
        'cx2': 'http://schemas.microsoft.com/office/drawing/2015/10/21/chartex',
        'cx3': 'http://schemas.microsoft.com/office/drawing/2016/5/9/chartex',
        'cx4': 'http://schemas.microsoft.com/office/drawing/2016/5/10/chartex',
        'cx5': 'http://schemas.microsoft.com/office/drawing/2016/5/11/chartex',
        'cx6': 'http://schemas.microsoft.com/office/drawing/2016/5/12/chartex',
        'cx7': 'http://schemas.microsoft.com/office/drawing/2016/5/13/chartex',
        'cx8': 'http://schemas.microsoft.com/office/drawing/2016/5/14/chartex',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'aink': 'http://schemas.microsoft.com/office/drawing/2016/ink',
        'am3d': 'http://schemas.microsoft.com/office/drawing/2017/model3d',
        'o': 'urn:schemas-microsoft-com:office:office',
        'oel': 'http://schemas.microsoft.com/office/2019/extlst',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'w10': 'urn:schemas-microsoft-com:office:word',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
        'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
        'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
        'w16du': 'http://schemas.microsoft.com/office/word/2023/wordml/word16du',
        'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
        'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wpi': 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk',
        'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    }

    # Criar o elemento raiz <w:document> com os namespaces
    root = ET.Element('w:document', attrib={
        'xmlns:wpc': namespaces['wpc'],
        'xmlns:cx': namespaces['cx'],
        'xmlns:cx1': namespaces['cx1'],
        'xmlns:cx2': namespaces['cx2'],
        'xmlns:cx3': namespaces['cx3'],
        'xmlns:cx4': namespaces['cx4'],
        'xmlns:cx5': namespaces['cx5'],
        'xmlns:cx6': namespaces['cx6'],
        'xmlns:cx7': namespaces['cx7'],
        'xmlns:cx8': namespaces['cx8'],
        'xmlns:mc': namespaces['mc'],
        'xmlns:aink': namespaces['aink'],
        'xmlns:am3d': namespaces['am3d'],
        'xmlns:o': namespaces['o'],
        'xmlns:oel': namespaces['oel'],
        'xmlns:r': namespaces['r'],
        'xmlns:m': namespaces['m'],
        'xmlns:v': namespaces['v'],
        'xmlns:wp14': namespaces['wp14'],
        'xmlns:wp': namespaces['wp'],
        'xmlns:w10': namespaces['w10'],
        'xmlns:w': namespaces['w'],
        'xmlns:w14': namespaces['w14'],
        'xmlns:w15': namespaces['w15'],
        'xmlns:w16cex': namespaces['w16cex'],
        'xmlns:w16cid': namespaces['w16cid'],
        'xmlns:w16': namespaces['w16'],
        'xmlns:w16du': namespaces['w16du'],
        'xmlns:w16sdtdh': namespaces['w16sdtdh'],
        'xmlns:w16se': namespaces['w16se'],
        'xmlns:wpg': namespaces['wpg'],
        'xmlns:wpi': namespaces['wpi'],
        'xmlns:wne': namespaces['wne'],
        'xmlns:wps': namespaces['wps'],
        'mc:Ignorable': 'w14 w15 w16se w16cid w16 w16cex w16sdtdh w16du wp14'
    })

    # Criar o corpo do documento
    body = ET.SubElement(root, f"{{{namespaces['w']}}}body")

    data = raw_json['Sheet1']
    print(f"Não conformidades: {len(data)}")
    
    # Ignorar a chave "ID"
    if 'ID' in data[0]:
        for item in data:
            item.pop('ID', None)

    # Iterar sobre os itens do JSON
    for item in data:
        for key, value in item.items():
            
            if key == "Fotos":
                print('inserinfo placeholder fotos _imagem_')
                p = ET.SubElement(body, f"{{{namespaces['w']}}}p")
                pPr = ET.SubElement(p, f"{{{namespaces['w']}}}pPr")
                ET.SubElement(pPr, f"{{{namespaces['w']}}}pStyle", {'w:val': 'a'})
                r = ET.SubElement(p, f"{{{namespaces['w']}}}r")
                ET.SubElement(r, f"{{{namespaces['w']}}}t").text = "_imagem_"
            
            else:
                text = value
                if text:
                    p = ET.SubElement(body, f"{{{namespaces['w']}}}p")
                    pPr = ET.SubElement(p, f"{{{namespaces['w']}}}pPr")
                    ET.SubElement(pPr, f"{{{namespaces['w']}}}pStyle", {'w:val': key})
                    r = ET.SubElement(p, f"{{{namespaces['w']}}}r")
                    ET.SubElement(r, f"{{{namespaces['w']}}}t").text = text
        
        # Adicionar quebra de página após cada item
        p = ET.SubElement(body, f"{{{namespaces['w']}}}p")
        r = ET.SubElement(p, f"{{{namespaces['w']}}}r")
        ET.SubElement(r, f"{{{namespaces['w']}}}br", {'w:type': 'page'})

    document_xml = ET.tostring(root, encoding='utf-8', method='xml').decode('utf-8')
    return document_xml


def replace_document_xml(docx_template_path, new_document_xml, output_docx_path):
    # Criar um diretório temporário
    temp_dir = 'temp_docx_dir'
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    # Extrair o conteúdo do arquivo .docx para o diretório temporário
    with zipfile.ZipFile(docx_template_path, 'r') as docx:
        docx.extractall(temp_dir)
    
    # Substituir o arquivo document.xml no diretório temporário
    document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
    with open(document_xml_path, 'w', encoding='utf-8') as file:
        file.write(new_document_xml)
    
    # Criar um novo arquivo .docx com o conteúdo do diretório temporário
    with zipfile.ZipFile(output_docx_path, 'w', zipfile.ZIP_DEFLATED) as docx:
        for root, dirs, files in os.walk(temp_dir):
            for file in files:
                file_path = os.path.join(root, file)
                # Calcula o caminho relativo para o arquivo dentro do ZIP
                relative_path = os.path.relpath(file_path, temp_dir)
                docx.write(file_path, relative_path)
    
    # Limpar o diretório temporário
    shutil.rmtree(temp_dir)
    
    print(f"O arquivo {output_docx_path} foi criado com sucesso.")




def inserir_imagens_no_word(word_path, output_docx_path, raw_json):
    # Abrir o documento Word
    document = Document(word_path)

    fotos_list = filtrar_imagens(raw_json['Sheet1'])
    print(len(fotos_list))
        
    image_count = 0
    # Iterar sobre os parágrafos do documento
    for paragraph in document.paragraphs:
        if '_imagem_' in paragraph.text:
            print("encontrado placeHolder _imagem_")
            # Para cada parágrafo com '_imagem_'
            print("fotos_list", len(fotos_list))
            for fotos_itens in fotos_list:
                print('fotos_itens', len(fotos_itens))
                
                # Substituir o texto '_imagem_'
                inline = paragraph.text.replace('_imagem_', '')
                paragraph.text = inline
                
                # Adicionar cada imagem da lista 'Fotos' no parágrafo
                for foto_base64 in fotos_itens:
                    # Decodificar a imagem base64
                    img_data = base64.b64decode(foto_base64)
                    
                    # Carregar a imagem usando Pillow
                    image = Image.open(BytesIO(img_data))
                    
                    # Definir uma altura fixa, mantendo a proporção
                    fixed_height = 200  # altura em pixels
                    width_percent = (fixed_height / float(image.size[1]))
                    new_width = int((float(image.size[0]) * float(width_percent)))
                    
                    # Redimensionar a imagem
                    resized_image = image.resize((new_width, fixed_height))
                    
                    # Converter a imagem redimensionada de volta para BytesIO
                    image_stream = BytesIO()
                    resized_image.save(image_stream, format=image.format)
                    image_stream.seek(0)  # Resetar o stream para a leitura

                    # Adicionar a imagem ao documento
                    run = paragraph.add_run()
                    run.add_picture(image_stream)

    # Salvar o documento modificado
    document.save(output_docx_path)
    
def filtrar_imagens(json_data):
    # Lista para armazenar arrays de imagens
    lista_de_arrays = []

    for item in json_data:
        for key, value in item.items():
            if key == 'Fotos':
                print(key)
                lista_de_arrays.append(value)
    return lista_de_arrays
